import streamlit as st
import re
from docx import Document
import PyPDF2
import pdfplumber
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
from io import BytesIO
from openai import OpenAI  # Updated import

# Initialize OpenAI client with API key from .streamlit/secrets.toml
try:
    api_key = st.secrets["OPENAI_API_KEY"]
    if not api_key:
        raise ValueError("API key is empty")
    
    client = OpenAI(api_key=api_key)  # Create client instance  
except Exception as e:
    st.error("❌ OpenAI API key not found! Please check your secrets configuration.")
    st.error(f"Error details: {str(e)}")
    st.info("""
    **Setup Instructions for Streamlit Cloud:**
    1. Go to your app's settings in Streamlit Cloud.
    2. Add to Secrets: OPENAI_API_KEY = "your_actual_openai_key_here"
    """)
    st.stop()


# -------------------------
# Image Preprocessing & OCR
# -------------------------
def preprocess_image_for_ocr(image):
    image = image.convert("L")  # Grayscale
    image = image.filter(ImageFilter.MedianFilter())
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(2)  # Increase contrast
    return image


def extract_text_with_ocr(image):
    try:
        processed_image = preprocess_image_for_ocr(image)
        text = pytesseract.image_to_string(processed_image, lang='eng')
        return text.strip()
    except Exception as e:
        st.error(f"OCR Error: {str(e)}")
        return ""


# -------------------------
# Text Detection Helpers
# -------------------------
def is_text_substantial(text, min_words=30):
    if not text or not text.strip():
        return False
    words = re.findall(r'\b\w+\b', text.lower())
    meaningful_words = [w for w in words if len(w) > 2]
    return len(meaningful_words) >= min_words


def find_index_pages(file, file_type):
    index_keywords = [
        'table of contents', 'contents', 'index', 'table of content',
        'chapter', 'section', 'outline', 'overview'
    ]
    potential_pages = []
    try:
        if file_type == 'pdf':
            file.seek(0)
            reader = PyPDF2.PdfReader(file)
            for page_num in range(min(10, len(reader.pages))):
                try:
                    page_text = reader.pages[page_num].extract_text().lower()
                    for keyword in index_keywords:
                        if keyword in page_text:
                            potential_pages.append(page_num)
                            break
                    if re.search(r'\d+\s*\.\s*\d+|\d+\s*-\s*\d+', page_text):
                        potential_pages.append(page_num)
                except Exception:
                    continue
        elif file_type == 'docx':
            file.seek(0)
            doc = Document(file)
            for para in doc.paragraphs[:50]:
                text = para.text.lower()
                for keyword in index_keywords:
                    if keyword in text:
                        potential_pages.append(0)
                        break
    except Exception as e:
        st.warning(f"Error while searching for index pages: {str(e)}")
    return list(set(potential_pages))


# -------------------------
# PDF/DOCX Text Extraction
# -------------------------
def extract_text_from_pdf_pages(file, page_numbers):
    combined_text = ""
    try:
        file.seek(0)
        reader = PyPDF2.PdfReader(file)
        for page_num in page_numbers:
            if page_num < len(reader.pages):
                page_text = reader.pages[page_num].extract_text() or ""
                if not page_text.strip():
                    try:
                        file.seek(0)
                        with pdfplumber.open(file) as pdf:
                            if page_num < len(pdf.pages):
                                page = pdf.pages[page_num]
                                page_text = page.extract_text() or ""
                                if not page_text.strip():
                                    pil_image = page.to_image(resolution=300).original
                                    page_text = extract_text_with_ocr(pil_image)
                    except Exception:
                        pass
                combined_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
    except Exception as e:
        st.warning(f"Error extracting text from multiple pages: {str(e)}")
    return combined_text.strip()


def extract_text_from_pdf(file):
    text = ""
    try:
        file.seek(0)
        reader = PyPDF2.PdfReader(file)
        if len(reader.pages) > 0:
            text = reader.pages[0].extract_text() or ""
    except Exception as e:
        st.warning(f"PyPDF2 extraction failed: {str(e)}")
    if not text.strip():
        try:
            file.seek(0)
            with pdfplumber.open(file) as pdf:
                if len(pdf.pages) > 0:
                    page = pdf.pages[0]
                    text = page.extract_text() or ""
                    if not text.strip():
                        pil_image = page.to_image(resolution=300).original
                        text = extract_text_with_ocr(pil_image)
        except Exception as e:
            st.warning(f"OCR extraction failed: {str(e)}")
    return text.strip()


def extract_text_from_docx(file):
    try:
        file.seek(0)
        doc = Document(file)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return full_text[:1500] if full_text else ""
    except Exception as e:
        st.error(f"DOCX extraction failed: {str(e)}")
        return ""


# -------------------------
# Smart Extraction Logic
# -------------------------
def extract_text_smart(file, file_type):
    extraction_log = []
    extraction_log.append("🔍 Attempting to extract text from first page...")
    if file_type == 'pdf':
        first_page_text = extract_text_from_pdf(file)
    else:
        first_page_text = extract_text_from_docx(file)
    if is_text_substantial(first_page_text):
        extraction_log.append("✅ First page contains substantial content")
        return first_page_text, extraction_log, "first_page"

    extraction_log.append("⚠️ First page has limited content, searching for index/table of contents...")
    index_pages = find_index_pages(file, file_type)
    if index_pages:
        extraction_log.append(f"📖 Found potential index pages: {[p + 1 for p in index_pages]}")
        if file_type == 'pdf':
            index_text = extract_text_from_pdf_pages(file, index_pages[:3])
        else:
            file.seek(0)
            doc = Document(file)
            full_text = "\n".join([para.text for para in doc.paragraphs[:100] if para.text.strip()])
            index_text = full_text[:3000]
        if is_text_substantial(index_text):
            extraction_log.append("✅ Successfully extracted content from index pages")
            return index_text, extraction_log, "index_pages"
        else:
            extraction_log.append("❌ Index pages also contain limited content")
    else:
        extraction_log.append("❌ No index/table of contents found")

    extraction_log.append("📄 Using first page content as fallback")
    return first_page_text, extraction_log, "first_page_fallback"


# -------------------------
# OpenAI Summarization (Updated for new API with debug info)
# -------------------------
def summarize_text_with_openai(text, extraction_method):
    # Debug info
    st.info("🤖 Making API call to OpenAI...")
    
    try:
        if extraction_method == "index_pages":
            system_prompt = """You are a helpful assistant that creates clear, concise summaries. 
            The text provided appears to be from a table of contents or index section. 
            Create a summary that captures the main topics and structure of the document based on this index information.
            Always format your response with each bullet point on a separate line using the format: - Bullet point text."""
            user_prompt = f"""Please analyze this table of contents/index and create a 4-6 bullet point summary:
            {text}"""
        else:
            system_prompt = """You are a helpful assistant that creates clear, concise summaries. 
            Always format your response with each bullet point on a separate line using the format: - Bullet point text."""
            user_prompt = f"""Please summarize the following text in 3-5 bullet points:
            {text}"""

        # Debug: Show which API endpoint is being called
        st.sidebar.write("🔄 Calling OpenAI API...")
        
        # Updated API call using the new client syntax
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=500,
            temperature=0.3
        )

        
        return response.choices[0].message.content.strip()

    except Exception as e:
        error_message = str(e)
        st.sidebar.error(f"❌ API Error: {error_message}")
        
        if "timeout" in error_message.lower():
            return "❌ Request timed out. Please try again."
        elif "rate limit" in error_message.lower():
            return "❌ Rate limit exceeded. Please wait a moment and try again."
        elif "insufficient_quota" in error_message.lower():
            return "❌ API quota exceeded. Please check your OpenAI account balance."
        elif "invalid_api_key" in error_message.lower():
            return "❌ Invalid API key. Please check your OpenAI API key configuration."
        else:
            return f"❌ API Error: {error_message}"


# -------------------------
# Streamlit UI
# -------------------------
def main():
    st.set_page_config(page_title="Smart Document Summarizer", page_icon="📄", layout="wide")
    st.title("📄 Smart Document Summarizer")
    st.markdown("*Enhanced OCR + AI-powered summarization with intelligent content detection*")

    with st.sidebar:
        st.header("📋 Instructions")
        st.markdown("""
        1. **Upload** a PDF or DOCX file
        2. **Review** the extracted text
        3. **Generate** an AI summary
        4. **Download** your summary
        """)
        st.header("⚙️ Features")
        st.markdown("""
        - ✅ OCR for scanned documents
        - ✅ Smart content detection
        - ✅ Index/TOC reading
        - ✅ AI summarization
        """)

    uploaded_file = st.file_uploader("Choose a file to summarize", type=["pdf", "docx"])

    if uploaded_file is not None:
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("File Name", uploaded_file.name)
        with col2:
            st.metric("File Size", f"{uploaded_file.size / 1024:.1f} KB")
        with col3:
            file_extension = uploaded_file.name.split(".")[-1].upper()
            st.metric("File Type", file_extension)

        with st.spinner("🔍 Analyzing document..."):
            extracted_text, extraction_log, extraction_method = extract_text_smart(uploaded_file, file_extension.lower())

        with st.expander("📋 Extraction Process Log", expanded=False):
            for log_entry in extraction_log:
                st.write(log_entry)

        if extracted_text:
            method_info = {
                "first_page": "✅ First Page Content",
                "index_pages": "📖 Table of Contents/Index",
                "first_page_fallback": "📄 First Page (Limited Content)"
            }
            st.info(f"**Content Source:** {method_info.get(extraction_method, 'Unknown')}")

            st.subheader("📜 Extracted Text")
            word_count = len(re.findall(r'\b\w+\b', extracted_text))
            st.info(f"📊 Extracted {len(extracted_text)} characters ({word_count} words)")

            with st.expander("🔍 View extracted text", expanded=False):
                st.text_area("Extracted Content", extracted_text, height=300, disabled=True)

            if st.button("✨ Generate AI Summary", type="primary", use_container_width=True):
                if len(extracted_text.strip()) < 20:
                    st.warning("⚠️ Text too short for summarization.")
                else:
                    with st.spinner("🤖 Generating AI summary..."):
                        summary = summarize_text_with_openai(extracted_text, extraction_method)

                    st.subheader("📝 AI-Generated Summary")
                    if summary and not summary.startswith("❌"):
                        bullet_markers = ['•', '-', '*']
                        formatted_summary = summary
                        for marker in bullet_markers:
                            formatted_summary = formatted_summary.replace(f'{marker} ', f'\n{marker} ')
                        lines = [line.strip() for line in formatted_summary.split('\n') if line.strip()]
                        formatted_summary = '\n'.join([f"{line}" for line in lines if line.startswith(tuple(bullet_markers))])
                        st.markdown(formatted_summary)
                    else:
                        st.markdown(summary)

                    if summary and not summary.startswith("❌"):
                        method_suffix = "_index" if extraction_method == "index_pages" else "_firstpage"
                        filename = f"summary_{uploaded_file.name.split('.')[0]}{method_suffix}.txt"
                        download_content = f"Document: {uploaded_file.name}\nContent Source: {method_info.get(extraction_method, 'Unknown')}\nSUMMARY:\n" + summary
                        st.download_button(
                            label="📥 Download Summary",
                            data=download_content,
                            file_name=filename,
                            mime="text/plain",
                            use_container_width=True
                        )
        else:
            st.error("❌ Unable to extract readable text. Try a different file.")


if __name__ == "__main__":
    main()
